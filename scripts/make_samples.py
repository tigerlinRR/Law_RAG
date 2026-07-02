#!/usr/bin/env python
"""Generate a few synthetic legal documents (docx + pdf) for testing the pipeline.
These are fictional and only exist to exercise ingestion and retrieval."""
from __future__ import annotations

from pathlib import Path

import docx
import fitz

OUT = Path(__file__).resolve().parent.parent / "data" / "sample"
OUT.mkdir(parents=True, exist_ok=True)


def make_docx(name: str, title: str, paragraphs: list[str]) -> None:
    d = docx.Document()
    d.add_heading(title, level=1)
    for p in paragraphs:
        d.add_paragraph(p)
    d.save(str(OUT / name))


def make_pdf(name: str, title: str, paragraphs: list[str]) -> None:
    doc = fitz.open()
    page = doc.new_page()
    text = title + "\n\n" + "\n\n".join(paragraphs)
    page.insert_textbox(fitz.Rect(50, 50, 545, 780), text, fontsize=11)
    doc.save(str(OUT / name))
    doc.close()


make_docx(
    "richtech_s8_2021.docx",
    "Form S-8 Registration Statement — Richtech Robotics Inc.",
    [
        "This Registration Statement on Form S-8 is filed by Richtech Robotics Inc. "
        "to register shares of common stock issuable under the company's 2021 Employee "
        "Stock Incentive Plan.",
        "The securities registered hereby consist of shares reserved for issuance to "
        "employees, directors, and consultants pursuant to stock options and restricted "
        "stock units granted under the Plan.",
        "The company incorporates by reference its Annual Report on Form 10-K and all "
        "subsequently filed periodic reports.",
    ],
)

make_docx(
    "acme_beta_nda_2022.docx",
    "Mutual Non-Disclosure Agreement",
    [
        "This Mutual Non-Disclosure Agreement is entered into between Acme Manufacturing "
        "LLC and Beta Components Inc. as of March 1, 2022.",
        "Each party agrees to hold the other party's Confidential Information in strict "
        "confidence and not to disclose it to any third party for a period of five years.",
        "Confidential Information does not include information that becomes publicly "
        "available through no breach of this Agreement.",
    ],
)

make_pdf(
    "richtech_equity_plan_2020.pdf",
    "Richtech Robotics Inc. — 2020 Equity Incentive Plan",
    [
        "The purpose of this Equity Incentive Plan is to attract and retain employees "
        "by providing equity-based compensation, including incentive stock options and "
        "restricted stock units.",
        "The maximum number of shares available for issuance under the Plan shall be "
        "subject to adjustment for stock splits, recapitalizations, and similar events.",
        "Vesting of awards shall generally occur over a four-year period, subject to the "
        "participant's continued service with the company.",
    ],
)

# A richer Master Services Agreement to exercise the due-diligence extractor.
# Includes a few deliberately one-sided / unusual terms to test risk flagging.
make_docx(
    "richtech_vendor_msa_2023.docx",
    "Master Services Agreement",
    [
        "This Master Services Agreement (the \"Agreement\") is entered into as of "
        "January 15, 2023 (the \"Effective Date\") by and between Richtech Robotics Inc. "
        "(\"Client\") and Nimbus Cloud Services LLC (\"Vendor\").",
        "1. Term. This Agreement shall commence on the Effective Date and continue for "
        "an initial term of three (3) years.",
        "2. Auto-Renewal. This Agreement shall automatically renew for successive "
        "one-year periods unless either party provides written notice of non-renewal at "
        "least ninety (90) days prior to the end of the then-current term.",
        "3. Termination. Client may terminate this Agreement for convenience upon sixty "
        "(60) days written notice. Vendor may terminate only for cause.",
        "4. Payment Terms. Client shall pay all undisputed invoices within thirty (30) "
        "days of receipt. Late payments accrue interest at 1.5% per month.",
        "5. Confidentiality. Each party shall protect the other's Confidential "
        "Information for a period of five (5) years following disclosure.",
        "6. Intellectual Property. All work product created by Vendor under this "
        "Agreement shall be the sole and exclusive property of Client upon full payment.",
        "7. Indemnification. Vendor shall indemnify and hold harmless Client against "
        "any third-party claims arising from Vendor's gross negligence or willful "
        "misconduct.",
        "8. Limitation of Liability. Notwithstanding anything to the contrary, Vendor's "
        "total liability under this Agreement shall not be limited with respect to "
        "breaches of confidentiality or indemnification obligations.",
        "9. Assignment. Vendor may assign this Agreement, including in connection with a "
        "change of control, without the prior consent of Client.",
        "10. Governing Law. This Agreement shall be governed by the laws of the State of "
        "Delaware, without regard to its conflict-of-laws principles.",
        "11. Dispute Resolution. Any dispute shall be resolved by binding arbitration "
        "administered by the American Arbitration Association in Wilmington, Delaware.",
    ],
)

# --- 8-K drafting experiment samples -----------------------------------------
# A synthetic HISTORICAL Item 1.01 8-K (precedent, for style/structure only) and a
# NEW supply contract (the drafting input) with a different counterparty/terms, so
# we can verify the drafted disclosure pulls facts from the new contract and never
# leaks the precedent's counterparty name, date, or dollar amount.
make_docx(
    "richtech_8k_2022_item101.docx",
    "Richtech Robotics Inc. — Form 8-K",
    [
        "Item 1.01 Entry into a Material Definitive Agreement.",
        "On June 3, 2022, Richtech Robotics Inc. (the \"Company\") entered into a "
        "Manufacturing Services Agreement (the \"Agreement\") with Orion Precision "
        "Manufacturing LLC (\"Orion\"), pursuant to which Orion will manufacture "
        "certain robotic subassemblies for the Company.",
        "The Agreement has an initial term of two (2) years and provides for "
        "aggregate minimum purchase commitments by the Company of approximately "
        "$4.2 million over the term. Either party may terminate the Agreement for "
        "an uncured material breach upon thirty (30) days' written notice.",
        "The foregoing description of the Agreement does not purport to be complete "
        "and is subject to the actual terms of the Agreement.",
    ],
)

make_docx(
    "richtech_vertex_supply_2024.docx",
    "Supply Agreement",
    [
        "This Supply Agreement (the \"Agreement\") is entered into as of September "
        "10, 2024 (the \"Effective Date\") by and between Richtech Robotics Inc. "
        "(\"Buyer\") and Vertex Robotics Components Inc. (\"Supplier\").",
        "1. Term. This Agreement shall commence on the Effective Date and continue "
        "for an initial term of three (3) years.",
        "2. Purchase Commitment. Buyer commits to purchase not less than $7.5 "
        "million in aggregate product volume over the initial term.",
        "3. Termination. Either party may terminate this Agreement upon a material "
        "breach that remains uncured for forty-five (45) days following written "
        "notice.",
        "4. Payment Terms. Buyer shall pay all undisputed invoices within forty-five "
        "(45) days of receipt.",
        "5. Exclusivity. Supplier agrees not to supply competing robotic components "
        "to any direct competitor of Buyer during the term of this Agreement.",
        "6. Governing Law. This Agreement shall be governed by the laws of the "
        "State of Delaware.",
    ],
)

print("Wrote sample docs to", OUT)
for p in sorted(OUT.iterdir()):
    print("  ", p.name)
