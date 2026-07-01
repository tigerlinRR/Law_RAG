"""Extract text from PDF and Word documents.

Returns a list of (page_number, text) blocks. For .docx there are no real pages,
so the whole document is returned as one block (page=None).

Scanned/image PDFs (no extractable text) are detected and flagged so the caller
knows OCR is needed. OCR itself is a Phase-2 add-on (kept out of the hot path so
the common case — electronic PDF/Word — stays fast and dependency-light).
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
import docx


@dataclass
class Block:
    page: int | None
    text: str


class NeedsOCR(Exception):
    """Raised when a PDF appears to be scanned images with no extractable text."""


def parse_pdf(path: Path, ocr_min_chars: int = 20) -> list[Block]:
    blocks: list[Block] = []
    total_chars = 0
    with fitz.open(path) as doc:
        for i, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            total_chars += len(text)
            if text:
                blocks.append(Block(page=i, text=text))
    if total_chars < ocr_min_chars:
        # Almost no text across the whole file -> it's a scanned/image PDF.
        raise NeedsOCR(f"{path.name}: no extractable text (likely scanned; OCR needed)")
    return blocks


def parse_docx(path: Path) -> list[Block]:
    d = docx.Document(str(path))
    parts: list[str] = [p.text for p in d.paragraphs if p.text.strip()]
    # Include table cell text — contracts often keep key terms in tables.
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return [Block(page=None, text="\n".join(parts))]


SUPPORTED = {".pdf", ".docx"}


def parse(path: Path) -> list[Block]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return parse_pdf(path)
    if ext == ".docx":
        return parse_docx(path)
    raise ValueError(f"Unsupported file type: {ext} ({path.name})")
