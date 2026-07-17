"""Export batch due-diligence reviews to Excel (comparison matrix) or Word (memo).

Excel is the classic DD "summary chart" — one row per contract, one column per
checklist clause, plus a risk count and a separate Risks sheet. Word is the
narrative memo form — full per-contract review with clause tables and risk lists.

Input `reviews` is a list of the dicts returned by summarize.review_contract().
"""
from __future__ import annotations

import io
import json
import re
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from .summarize import CHECKLIST

_HDR_FILL = PatternFill("solid", fgColor="1A2238")
_HDR_FONT = Font(bold=True, color="FFFFFF")


def _clause_map(review: dict) -> dict:
    return {c.get("name", ""): c.get("value", "") for c in review.get("clauses", [])}


def to_excel(reviews: list[dict]) -> bytes:
    wb = Workbook()

    ws = wb.active
    ws.title = "Summary"
    headers = ["File", "Type", "Parties"] + CHECKLIST + ["# Risks"]
    ws.append(headers)
    for review in reviews:
        cmap = _clause_map(review)
        row = [review.get("_source", ""), review.get("doc_type", ""),
               "; ".join(review.get("parties", []))]
        row += [cmap.get(clause, "") for clause in CHECKLIST]
        row.append(len(review.get("key_risks", [])))
        ws.append(row)

    risks = wb.create_sheet("Risks")
    risks.append(["File", "Risk"])
    for review in reviews:
        for r in review.get("key_risks", []):
            risks.append([review.get("_source", ""), r])

    for sheet in (ws, risks):
        for col, name in enumerate(sheet[1], start=1):
            name.fill = _HDR_FILL
            name.font = _HDR_FONT
            name.alignment = Alignment(vertical="top", wrap_text=True)
            width = 16 if col > 1 else 28
            sheet.column_dimensions[get_column_letter(col)].width = width
        sheet.freeze_panes = "A2"
        for r in sheet.iter_rows(min_row=2):
            for cell in r:
                cell.alignment = Alignment(vertical="top", wrap_text=True)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def to_word(reviews: list[dict]) -> bytes:
    doc = docx.Document()
    doc.add_heading("Due Diligence Report", 0)
    intro = doc.add_paragraph(
        f"{len(reviews)} document(s) reviewed. AI-assisted summary — verify each "
        "finding against the source document before relying on it.")
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for i, review in enumerate(reviews):
        title = review.get("_source", "Contract")
        if review.get("doc_type"):
            title += f"  ({review['doc_type']})"
        doc.add_heading(title, level=1)

        if review.get("summary"):
            doc.add_paragraph(review["summary"])
        if review.get("parties"):
            doc.add_paragraph("Parties: " + "; ".join(review["parties"]))

        clauses = review.get("clauses", [])
        if clauses:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Clause", "Value", "Source quote"
            for c in clauses:
                cells = table.add_row().cells
                cells[0].text = c.get("name", "")
                cells[1].text = c.get("value", "")
                cells[2].text = c.get("quote", "")

        risks = review.get("key_risks", [])
        if risks:
            doc.add_heading("Key risks to review", level=2)
            for r in risks:
                doc.add_paragraph(r, style="List Bullet")

        if i < len(reviews) - 1:
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------- 8-K draft export (experimental) ----------
#
# Mirrors the actual SEC Form 8-K structure (cover page -> Item disclosure ->
# Item 9.01 exhibit index -> signature) rather than a generic report layout, so
# a lawyer sees something recognizable as a filing draft. Cover-page fields
# below (registrant/EIN/address/securities) are Richtech-specific -- this tool
# currently drafts for one registrant; parameterize per-client if that changes.
# Our own transparency material (precedents, fact trace) is kept as a clearly
# separate appendix, never mixed into the filing-shaped section.

# Built-in fallback registrant profile. NOT hardcoded into the filing path anymore — the
# active profile comes from load_registrant() (registrant.json), so a different issuer, a
# changed address, or new officers are a config edit, not a code change. This default keeps
# the current deployment working if no registrant.json is present.
_DEFAULT_REGISTRANT = {
    "name": "Richtech Robotics Inc.",
    "state": "Nevada",
    "file_number": "001-41866",
    "irs_ein": "88-2870106",
    "address": ["2975 Lincoln Rd,", "Las Vegas, NV 89115"],
    "phone": "(866) 236-3835",
    "securities": [("Class B Common Stock, par value $0.0001 per share", "RR",
                     "The Nasdaq Stock Market LLC")],
    "emerging_growth_company": True,
    "signer_name": "Zhenwu (Wayne) Huang",
    "signer_title": "Chief Executive Officer and Director",
}


def load_registrant() -> dict:
    """The registrant profile used on the cover page + signature block. Loaded from the
    configured registrant.json (per-deployment), merged over the built-in default so a partial
    file still works. Makes the tool multi-company and lets the address/officers be updated
    without touching code. `securities` rows are normalized to tuples for the table renderers."""
    from .config import CONFIG
    data = dict(_DEFAULT_REGISTRANT)
    p = Path(CONFIG.registrant_file)
    if p.exists():
        try:
            data.update(json.loads(p.read_text(encoding="utf-8")))
        except Exception:
            pass  # malformed file -> fall back to the default rather than crash a filing
    data["securities"] = [tuple(s) for s in data.get("securities", [])]
    return data


REGISTRANT = load_registrant()


def _report_date(draft: dict) -> str:
    """The 8-K 'date of earliest event reported' — the transaction date, which
    is almost always the first date stated in the disclosure ('On <date>, ...
    entered into ...'). Fall back to any date in the cited facts."""
    m = re.search(r"\b([A-Z][a-z]+ \d{1,2}, \d{4})\b", draft.get("disclosure", ""))
    if m:
        return m.group(1)
    for f in draft.get("facts_used") or []:
        m = re.search(r"\b([A-Z][a-z]+ \d{1,2}, \d{4})\b", f.get("fact", ""))
        if m:
            return m.group(1)
    return "[DATE]"


def filing_date_iso(draft: dict) -> str:
    """The contract/event date in YYYY-MM-DD, for naming downloaded files after the
    actual transaction date (matches how the rest of the project's own historical
    filings/contracts are named) rather than an opaque internal generation id."""
    import datetime
    try:
        return datetime.datetime.strptime(_report_date(draft), "%B %d, %Y").strftime("%Y-%m-%d")
    except ValueError:
        return "undated"


def _split_paragraphs(disclosure: str, item: str) -> list[str]:
    """Split a disclosure string into paragraphs, dropping a leading 'Item X.XX...'
    line if the model repeated the heading we already render separately."""
    paras = [p.strip() for p in (disclosure or "").split("\n\n") if p.strip()]
    if paras and item and paras[0].lower().startswith(f"item {item}".lower()):
        paras = paras[1:]
    return paras


def _disclosure_paragraphs(draft: dict) -> list[str]:
    return _split_paragraphs(draft.get("disclosure") or "", str(draft.get("item", "")))


def _filing_sections(draft: dict) -> list[dict]:
    """Ordered Item sections that make up the filing body. Uses `_items` (multi-Item)
    when present; falls back to the single top-level item for older records."""
    secs = draft.get("_items")
    if secs:
        return secs
    return [{"item": draft.get("item", ""), "item_title": draft.get("item_title", ""),
             "disclosure": draft.get("disclosure", ""), "cross_ref": False}]


def _rule_para(doc, before=0, after=8):
    """A thick horizontal rule (EDGAR brackets each page's content top and bottom)."""
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def _body_para(doc, text: str, *, justify=True, indent=True, bold_terms=True):
    """A justified, first-line-indented body paragraph with defined terms bolded,
    matching the real filing's Item text."""
    from docx.shared import Inches
    p = doc.add_paragraph()
    if justify:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.4)
    if bold_terms:
        for seg, is_bold in _bold_segments(text):
            run = p.add_run(seg)
            run.bold = is_bold
    else:
        p.add_run(text)
    return p


def draft_to_word(draft: dict) -> bytes:
    from docx.shared import Inches, Pt
    r = REGISTRANT
    date = _report_date(draft)
    doc = docx.Document()

    # Compact spacing + modest margins so the whole cover page fits on ONE page,
    # like a real filing (spacers pushed it onto a second page otherwise).
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"  # match the real filing's serif face
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.5)
        s.left_margin = s.right_margin = Inches(0.85)

    def centered(text: str, bold=False, size=None, gap_before=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if gap_before:
            p.paragraph_format.space_before = Pt(gap_before)
        run = p.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        return p

    _rule_para(doc, after=8)
    centered("UNITED STATES", bold=True)
    centered("SECURITIES AND EXCHANGE COMMISSION", bold=True)
    centered("Washington, D.C. 20549", bold=True)
    centered("FORM 8-K", bold=True, size=14, gap_before=7)
    centered("CURRENT REPORT", bold=True, gap_before=7)
    centered("PURSUANT TO SECTION 13 OR 15(d) OF THE", bold=True)
    centered("SECURITIES EXCHANGE ACT OF 1934", bold=True)
    centered(f"Date of Report (Date of earliest event reported): {date}", gap_before=7)
    centered(r["name"], bold=True, gap_before=7)
    centered("(Exact name of registrant as specified in its charter)")

    # Registrant id: values underlined (bottom border), labels below — no grid box.
    t = doc.add_table(rows=2, cols=3)
    _no_table_borders(t)
    for i, v in enumerate([r["state"], r["file_number"], r["irs_ein"]]):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].add_run(v).bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _cell_borders(cell, bottom=True)
    for i, v in enumerate(["(State or other jurisdiction of incorporation)",
                            "(Commission File Number)", "(IRS Employer Identification No.)"]):
        p = t.rows[1].cells[i].paragraphs[0]
        p.add_run(v).font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for j, line in enumerate(r["address"]):
        centered(line, bold=True, gap_before=6 if j == 0 else 0)
    centered("(Address of principal executive offices, including zip code)")
    centered(f"Registrant's telephone number, including area code: {r['phone']}", gap_before=5)
    centered("Not Applicable", bold=True, gap_before=5)
    centered("(Former name or former address, if changed since last report)")

    intro = _body_para(doc,
        "Check the appropriate box below if the Form 8-K filing is intended to "
        "simultaneously satisfy the filing obligation of the registrant under any of "
        "the following provisions:", bold_terms=False)
    intro.paragraph_format.space_before = Pt(6)
    doc.add_paragraph("☐  Written communications pursuant to Rule 425 under the Securities Act (17 CFR 230.425)")
    doc.add_paragraph("☐  Soliciting material pursuant to Rule 14a-12 under the Exchange Act (17 CFR 240.14a-12)")
    doc.add_paragraph("☐  Pre-commencement communications pursuant to Rule 14d-2(b) under the Exchange Act (17 CFR 240.14d-2(b))")
    doc.add_paragraph("☐  Pre-commencement communications pursuant to Rule 13e-4(c) under the Exchange Act (17 CFR 240.13e-4(c))")

    centered("Securities registered pursuant to Section 12(b) of the Act:", gap_before=5)
    # Securities table: underlined header (top+bottom border) + shaded data row, no grid.
    t2 = doc.add_table(rows=1 + len(r["securities"]), cols=3)
    _no_table_borders(t2)
    for i, h in enumerate(["Title of each class", "Trading Symbol(s)", "Name of each exchange on which registered"]):
        cell = t2.rows[0].cells[i]
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _cell_borders(cell, top=True, bottom=True)
    for row_i, (cls, sym, exch) in enumerate(r["securities"], start=1):
        for i, v in enumerate((cls, sym, exch)):
            cell = t2.rows[row_i].cells[i]
            cell.text = v
            if i:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _shade_cell(cell)
            _cell_borders(cell, bottom=True)

    egc = "☒" if r["emerging_growth_company"] else "☐"
    egp = _body_para(doc,
        "Indicate by check mark whether the registrant is an emerging growth company as "
        "defined in Rule 405 of the Securities Act of 1933 (§230.405 of this chapter) or "
        "Rule 12b-2 of the Securities Exchange Act of 1934 (§240.12b-2 of this chapter).",
        bold_terms=False)
    egp.paragraph_format.space_before = Pt(6)
    doc.add_paragraph(f"Emerging growth company {egc}")
    _body_para(doc,
        "If an emerging growth company, indicate by check mark if the registrant has "
        "elected not to use the extended transition period for complying with any new or "
        "revised financial accounting standards provided pursuant to Section 13(a) of the "
        "Exchange Act. ☐", bold_terms=False)
    _rule_para(doc, before=8, after=0)

    doc.add_page_break()

    _rule_para(doc, after=8)
    fls = draft.get("_forward_looking_statements")
    primary = str(draft.get("item", ""))  # FLS belongs right after the item it relates to
    for sec in _filing_sections(draft):
        hdr = doc.add_paragraph()
        hdr.paragraph_format.space_before = Pt(6)
        hdr.add_run(f"Item {sec.get('item','')}. {sec.get('item_title','')}.").bold = True
        for para in _split_paragraphs(sec.get("disclosure", ""), str(sec.get("item", ""))):
            _body_para(doc, para)
        if fls and str(sec.get("item", "")) == primary:
            fp = doc.add_paragraph()
            fp.paragraph_format.space_before = Pt(6)
            fp.add_run("Forward-Looking Statements").bold = True
            _body_para(doc, fls, indent=False, bold_terms=False)

    doc.add_paragraph().add_run("Item 9.01. Financial Statements and Exhibits.").bold = True
    ital = doc.add_paragraph()
    ital.add_run("(d) Exhibits").italic = True
    doc.add_paragraph("The following exhibits are being filed herewith:")
    # Exhibit index: underlined header, no grid.
    t3 = doc.add_table(rows=3, cols=2)
    _no_table_borders(t3)
    for i, h in enumerate(["Exhibit No.", "Description"]):
        cell = t3.rows[0].cells[i]
        cell.paragraphs[0].add_run(h).bold = True
        _cell_borders(cell, bottom=True)
    t3.rows[1].cells[0].text = "10.1"
    t3.rows[1].cells[1].text = f"{draft.get('_doc_type') or 'Agreement'}, dated {date}"
    t3.rows[2].cells[0].text = "104"
    t3.rows[2].cells[1].text = "Cover Page Interactive Data File (embedded within the Inline XBRL document)"
    _rule_para(doc, before=10, after=2)
    centered("1")

    doc.add_page_break()
    _rule_para(doc, after=8)
    centered("SIGNATURE", bold=True)
    _body_para(doc,
        "Pursuant to the requirements of the Securities Exchange Act of 1934, the "
        "registrant has duly caused this report to be signed on its behalf by the "
        "undersigned hereunto duly authorized.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(r["name"]).bold = True
    doc.add_paragraph()
    p = doc.add_paragraph(f"By: /s/ {r['signer_name']}  [DRAFT — NOT YET SIGNED]")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph(f"Name: {r['signer_name']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph(f"Title: {r['signer_title']}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f"Dated: {date}")
    _rule_para(doc, before=10, after=2)
    centered("2")

    # NOTE: the filing document ends here — no review/QC material is mixed in, so
    # this file is the clean 8-K ready for counsel to finalize. The review pack
    # (precedents, fact trace, SEC checks, full extraction) is a SEPARATE document,
    # produced by review_to_word() / review_to_pdf().
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_GUARDRAIL_VERDICT = {
    "blocked": "BLOCKED — the draft contains figures not grounded in the source "
               "contract. Do not treat as ready; resolve every RED item below.",
    "needs_review": "NEEDS REVIEW (does not block) — the draft contains figures that are "
                    "DERIVED from the source (confirm the arithmetic shown) and/or required "
                    "source figures absent from the draft. Review each below.",
    "clean": "CLEAN — every figure in the draft is grounded in the source contract.",
}


def review_to_word(draft: dict) -> bytes:
    """The review pack — a SEPARATE document (never mixed into the filing) that
    lets counsel check the draft: SEC-requirement checks, precedents used, the
    fact -> source-quote trace, and the full set of extracted contract terms."""
    doc = docx.Document()
    doc.add_heading(f"8-K Draft — Review Materials (Item {draft.get('item', '')})", 0)
    doc.add_paragraph(
        f"Companion to the drafted 8-K for source contract "
        f"{draft.get('_source_contract', '—')}. NOT part of the filing — for "
        "internal review only. Verify every fact below against the source contract.")

    guard = draft.get("_guardrail") or {}
    if guard.get("items") is not None:
        verdict = guard.get("verdict", "")
        doc.add_heading("Fact reconciliation (source grounding)", level=1)
        doc.add_paragraph(_GUARDRAIL_VERDICT.get(verdict, verdict))
        gitems = guard.get("items") or []
        fabricated = [i for i in gitems if i.get("status") == "fabricated"]
        derived = [i for i in gitems if i.get("status") == "derived"]
        omitted = [i for i in gitems if i.get("status") == "omitted"]
        if fabricated:
            doc.add_paragraph(
                "RED — figures in the draft NOT found in the source contract "
                "(likely fabricated; do not file until resolved):")
            for i in fabricated:
                doc.add_paragraph(f"{i.get('kind','')}: {i.get('raw','')}",
                                  style="List Bullet")
        if derived:
            doc.add_paragraph(
                "DERIVED — figures computed from verbatim source figures (grounded, but "
                "confirm the arithmetic):")
            for i in derived:
                doc.add_paragraph(f"{i.get('kind','')}: {i.get('raw','')}  "
                                  f"{i.get('source_snippet','')}", style="List Bullet")
        if omitted:
            doc.add_paragraph(
                "AMBER — figures in the source contract absent from the draft. 8-K "
                "disclosure is selective; confirm none of these is legally material:")
            for i in omitted:
                doc.add_paragraph(f"{i.get('kind','')}: {i.get('raw','')}",
                                  style="List Bullet")

    compliance = draft.get("_compliance") or []
    if compliance:
        doc.add_heading("SEC requirement checks", level=1)
        for c in compliance:
            mark = "✓" if c.get("satisfied") else "✗ MISSING"
            doc.add_paragraph(f"{mark}  {c.get('requirement', '')}", style="List Bullet")

    precedents = draft.get("_precedents_used") or []
    if precedents:
        doc.add_heading("Precedents used (style reference only)", level=1)
        for p in precedents:
            doc.add_paragraph(p, style="List Bullet")

    facts = draft.get("facts_used") or []
    if facts:
        doc.add_heading("Fact -> source trace", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        fhdr = table.rows[0].cells
        fhdr[0].text, fhdr[1].text, fhdr[2].text = "Fact", "Source quote", "Verified"
        for f in facts:
            cells = table.add_row().cells
            cells[0].text = f.get("fact", "")
            cells[1].text = f.get("source_quote", "")
            if f.get("source") == "business_context":
                cells[2].text = "Business input (not a contract citation — reviewer-provided)"
            else:
                cells[2].text = "Yes" if f.get("verified") else "⚠ UNVERIFIED"

    all_terms = draft.get("_all_extracted_terms") or []
    if all_terms:
        doc.add_heading("All terms extracted from the contract", level=1)
        doc.add_paragraph(
            "The disclosure states only the material terms, following 8-K convention. "
            "This is the full set the review engine extracted — use it to confirm "
            "nothing material was left out of the disclosure.")
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        thdr = table.rows[0].cells
        thdr[0].text, thdr[1].text = "Term", "Value"
        for t in all_terms:
            cells = table.add_row().cells
            cells[0].text = t.get("name", "")
            cells[1].text = t.get("value", "")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _esc(s: str) -> str:
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# Real 8-K filings set off a newly-defined term in bold the first time it appears,
# e.g. (the "Company"), ("Purchase and Sale Agreement"), (the "Earnest Money").
# We bold ONLY quoted terms that sit inside a parenthetical, matching that
# convention (a later plain "the Company" is not bold, and quoted phrases outside
# parentheses — like "forward-looking statements" in the FLS legend — are not).
_PAREN_RE = re.compile(r"\([^)]*\)")
_QUOTED_RE = re.compile(r"[“\"][^”\"]+[”\"]")


def _bold_segments(text: str) -> list[tuple[str, bool]]:
    """Split `text` into (segment, is_bold) runs, bolding defined terms."""
    spans: list[tuple[int, int]] = []
    for pm in _PAREN_RE.finditer(text):
        for qm in _QUOTED_RE.finditer(pm.group(0)):
            spans.append((pm.start() + qm.start(), pm.start() + qm.end()))
    if not spans:
        return [(text, False)]
    segs: list[tuple[str, bool]] = []
    i = 0
    for s, e in spans:
        if s > i:
            segs.append((text[i:s], False))
        segs.append((text[s:e], True))
        i = e
    if i < len(text):
        segs.append((text[i:], False))
    return segs


def _bold_defined_html(text: str) -> str:
    """HTML with defined terms bolded (input is raw text; each piece is escaped)."""
    return "".join(f"<b>{_esc(t)}</b>" if b else _esc(t) for t, b in _bold_segments(text))


def _cell_borders(cell, **sides) -> None:
    """Set individual cell borders (Word). sides e.g. bottom=True, top=True — real
    8-K tables use underlines (a bottom border), not a full grid box."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side, on in sides.items():
        if not on:
            continue
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "000000")
        tcB.append(el)
    tcPr.append(tcB)


def _shade_cell(cell, fill: str = "CFE2F3") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _no_table_borders(table) -> None:
    """Remove all borders from a table (python-docx default table already has none,
    but be explicit so no theme grid sneaks in)."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tblPr.append(borders)


def _draft_html(draft: dict) -> str:
    """The clean filing document — cover page, Item disclosure, Item 9.01, and
    signature. NO review/QC material (that lives in _review_html)."""
    esc = _esc
    r = REGISTRANT
    date = _report_date(draft)

    fls = draft.get("_forward_looking_statements")
    primary = str(draft.get("item", ""))  # FLS belongs right after the item it relates to

    def _sec_html(sec):
        html = (f"<p class='heading'>Item {_esc(sec.get('item',''))}. "
                f"{_esc(sec.get('item_title',''))}.</p>"
                + "".join(f"<p class='body'>{_bold_defined_html(para)}</p>"
                          for para in _split_paragraphs(sec.get('disclosure', ''),
                                                        str(sec.get('item', '')))))
        if fls and str(sec.get('item', '')) == primary:
            html += (f"<p class='heading'>Forward-Looking Statements</p>"
                     f"<p class='body' style='text-indent:0;'>{esc(fls)}</p>")
        return html

    disclosure_html = "".join(_sec_html(sec) for sec in _filing_sections(draft))
    fls_html = ""  # inlined after the primary Item above (not appended at the end)
    sec_rows = "".join(
        f"<tr><td class='left'>{esc(cls)}</td><td>{esc(sym)}</td><td>{esc(exch)}</td></tr>"
        for cls, sym, exch in r["securities"]
    )
    egc = "&#9746;" if r["emerging_growth_company"] else "&#9744;"  # ☒ / ☐

    return f"""<!doctype html><html><head><meta charset="utf-8"><style>
      body {{ font-family: 'Times New Roman', Georgia, serif; color: #000;
              font-size: 12px; line-height: 1.32; margin: 0; }}
      .page {{ padding: 0; }}
      .center {{ text-align: center; }}
      .bold {{ font-weight: 700; }}
      .small {{ font-size: 10.5px; }}
      /* EDGAR-style thick rule bracketing each page's content, top and bottom. */
      .pgrule {{ border: none; border-top: 2px solid #000; margin: 0 0 10px; }}
      .pgrule.bottom {{ margin: 12px 0 0; }}
      .pageno {{ text-align: center; font-size: 11px; margin-top: 4px; }}
      .pagebreak {{ page-break-before: always; }}

      .cover p {{ margin: 2px 0; text-align: center; }}
      .cover .gap {{ margin-top: 10px; }}
      /* Registrant id row: values underlined (bottom border), labels below — no box. */
      table.cover-id {{ width: 100%; border-collapse: collapse; margin: 8px 0 2px; }}
      table.cover-id td {{ text-align: center; padding: 2px 8px; vertical-align: top; }}
      table.cover-id .vals td {{ font-weight: 700; border-bottom: 1px solid #000; }}
      table.cover-id .labels td {{ font-size: 10.5px; }}
      /* Securities table: underlined header + shaded data row, no vertical lines. */
      table.sec {{ width: 100%; border-collapse: collapse; margin: 6px 0; font-size: 11px; }}
      table.sec th {{ border-top: 1px solid #000; border-bottom: 1px solid #000;
        font-weight: 700; text-align: center; padding: 4px 6px; }}
      table.sec td {{ background: #cfe2f3; text-align: center; padding: 4px 6px;
        border-bottom: 1px solid #000; }}
      table.sec td.left {{ text-align: left; }}
      .cover p.boxnote {{ text-indent: 0.4in; margin: 8px 0; text-align: justify; }}
      .cover p.checkline {{ margin: 5px 0 5px 0.3in; text-align: left; }}

      .heading {{ font-weight: 700; margin: 16px 0 8px; }}
      p.body {{ text-indent: 0.4in; margin: 0 0 10px; text-align: justify; }}
      /* Exhibit index: underlined header, no grid. */
      table.exhibit-idx {{ width: 100%; border-collapse: collapse; margin: 4px 0; }}
      table.exhibit-idx th {{ border-bottom: 1px solid #000; text-align: left;
        font-weight: 700; padding: 3px 6px; }}
      table.exhibit-idx td {{ padding: 3px 6px; vertical-align: top; }}
      table.exhibit-idx td.exno {{ width: 12%; white-space: nowrap; }}
      .sig-block {{ width: 62%; margin-left: auto; margin-top: 34px; }}
      .sig-block p {{ margin: 3px 0; }}
      .sig-line {{ border-bottom: 1px solid #000; }}
    </style></head><body>

      <div class="page cover">
        <hr class="pgrule">
        <p class="bold">UNITED STATES</p>
        <p class="bold">SECURITIES AND EXCHANGE COMMISSION</p>
        <p class="bold">Washington, D.C. 20549</p>
        <p class="bold gap" style="font-size:15px;">FORM 8-K</p>
        <p class="bold gap">CURRENT REPORT</p>
        <p class="bold gap">PURSUANT TO SECTION 13 OR 15(d) OF THE<br>SECURITIES EXCHANGE ACT OF 1934</p>
        <p class="gap">Date of Report (Date of earliest event reported): <span class="bold">{esc(date)}</span></p>
        <p class="bold gap">{esc(r['name'])}</p>
        <p class="small">(Exact name of registrant as specified in its charter)</p>
        <table class="cover-id">
          <tr class="vals"><td>{esc(r['state'])}</td><td>{esc(r['file_number'])}</td><td>{esc(r['irs_ein'])}</td></tr>
          <tr class="labels"><td>(State or other jurisdiction<br>of incorporation)</td>
            <td>(Commission File Number)</td><td>(IRS Employer<br>Identification No.)</td></tr>
        </table>
        <p class="bold gap">{"<br>".join(esc(l) for l in r['address'])}</p>
        <p class="small">(Address of principal executive offices, including zip code)</p>
        <p class="gap">Registrant's telephone number, including area code: <span class="bold">{esc(r['phone'])}</span></p>
        <p class="bold gap">Not Applicable</p>
        <p class="small">(Former name or former address, if changed since last report)</p>

        <p class="boxnote">Check the appropriate box below if the Form 8-K filing is intended to
        simultaneously satisfy the filing obligation of the registrant under any of the following provisions:</p>
        <p class="checkline">&#9744;&nbsp; Written communications pursuant to Rule 425 under the Securities Act (17 CFR 230.425)</p>
        <p class="checkline">&#9744;&nbsp; Soliciting material pursuant to Rule 14a-12 under the Exchange Act (17 CFR 240.14a-12)</p>
        <p class="checkline">&#9744;&nbsp; Pre-commencement communications pursuant to Rule 14d-2(b) under the Exchange Act (17 CFR 240.14d-2(b))</p>
        <p class="checkline">&#9744;&nbsp; Pre-commencement communications pursuant to Rule 13e-4(c) under the Exchange Act (17 CFR 240.13e-4(c))</p>

        <p class="center" style="margin-top:8px;">Securities registered pursuant to Section 12(b) of the Act:</p>
        <table class="sec">
          <thead><tr><th>Title of each class</th><th>Trading Symbol(s)</th>
          <th>Name of each exchange on which registered</th></tr></thead>
          <tbody>{sec_rows}</tbody>
        </table>

        <p class="boxnote">Indicate by check mark whether the registrant is an emerging growth company as
        defined in Rule 405 of the Securities Act of 1933 (&sect;230.405 of this chapter) or Rule 12b-2 of the
        Securities Exchange Act of 1934 (&sect;240.12b-2 of this chapter).</p>
        <p class="checkline">Emerging growth company {egc}</p>
        <p class="boxnote">If an emerging growth company, indicate by check mark if the registrant has elected not
        to use the extended transition period for complying with any new or revised financial accounting
        standards provided pursuant to Section 13(a) of the Exchange Act. &#9744;</p>
        <hr class="pgrule bottom">
      </div>

      <div class="pagebreak page">
        <hr class="pgrule">
        {disclosure_html}
        {fls_html}
        <p class="heading">Item 9.01. Financial Statements and Exhibits.</p>
        <p class="body" style="text-indent:0; font-style:italic; margin-bottom:4px;">(d) Exhibits</p>
        <p class="body" style="text-indent:0; margin-bottom:4px;">The following exhibits are being filed herewith:</p>
        <table class="exhibit-idx">
          <thead><tr><th>Exhibit No.</th><th>Description</th></tr></thead>
          <tbody>
            <tr><td class="exno">10.1</td><td>{esc(draft.get('_doc_type') or 'Agreement')}, dated {esc(date)}</td></tr>
            <tr><td class="exno">104</td><td>Cover Page Interactive Data File (embedded within the Inline XBRL document)</td></tr>
          </tbody>
        </table>
        <hr class="pgrule bottom">
        <div class="pageno">1</div>
      </div>

      <div class="pagebreak page">
        <hr class="pgrule">
        <p class="center bold" style="margin-bottom:12px;">SIGNATURE</p>
        <p class="body">Pursuant to the requirements of the Securities Exchange Act of 1934, the registrant has
        duly caused this report to be signed on its behalf by the undersigned hereunto duly authorized.</p>
        <div class="sig-block">
          <p class="bold">{esc(r['name'])}</p>
          <p style="margin-top:24px;"><span class="sig-line">By: /s/ {esc(r['signer_name'])}</span>
             &nbsp; <i>[DRAFT — NOT YET SIGNED]</i></p>
          <p>Name: {esc(r['signer_name'])}</p>
          <p>Title: {esc(r['signer_title'])}</p>
        </div>
        <p style="margin-top:24px;">Dated: {esc(date)}</p>
        <hr class="pgrule bottom">
        <div class="pageno">2</div>
      </div>
    </body></html>"""


def _review_html(draft: dict) -> str:
    """The review pack as a SEPARATE document — SEC checks, precedents, fact
    trace, full extraction. Never mixed into the filing."""
    esc = _esc
    def _fact_row(f):
        if f.get("source") == "business_context":
            status = "Business input (not a contract citation — reviewer-provided)"
            cls = ""
        else:
            status = "Yes" if f.get("verified") else "⚠ UNVERIFIED"
            cls = "" if f.get("verified") else ' class="unverified"'
        return (f"<tr{cls}><td>{esc(f.get('fact',''))}</td>"
                f"<td>{esc(f.get('source_quote',''))}</td><td>{esc(status)}</td></tr>")
    facts_rows = "".join(_fact_row(f) for f in draft.get("facts_used") or [])
    precedents = "".join(f"<li>{esc(p)}</li>" for p in draft.get("_precedents_used") or [])
    all_terms_rows = "".join(
        f"<tr><td>{esc(t.get('name',''))}</td><td>{esc(t.get('value',''))}</td></tr>"
        for t in draft.get("_all_extracted_terms") or []
    )
    compliance_rows = "".join(
        f"<tr><td>{'✓' if c.get('satisfied') else '✗ MISSING'}</td>"
        f"<td>{esc(c.get('requirement',''))}</td></tr>"
        for c in draft.get("_compliance") or []
    )
    guard = draft.get("_guardrail") or {}
    guard_html = ""
    if guard.get("items") is not None:
        gitems = guard.get("items") or []
        def _grows(status, show_snip):
            return "".join(
                f"<tr><td>{esc(i.get('kind',''))}</td><td>{esc(i.get('raw',''))}</td>"
                f"<td>{esc(i.get('source_snippet') or '') if show_snip else ''}</td></tr>"
                for i in gitems if i.get("status") == status)
        fab = _grows("fabricated", False)
        der = _grows("derived", True)
        om = _grows("omitted", True)
        verdict = guard.get("verdict", "")
        cls = {"blocked": "gr-red", "needs_review": "gr-amber"}.get(verdict, "gr-green")
        guard_html = (
            f"<h2>Fact reconciliation (source grounding)</h2>"
            f"<p class='verdict {cls}'>{esc(_GUARDRAIL_VERDICT.get(verdict, verdict))}</p>"
            + ("<p class='note'>RED — in the draft, NOT found in the source contract "
               "(likely fabricated):</p><table><thead><tr><th>Kind</th><th>Draft value"
               "</th><th></th></tr></thead><tbody>" + fab + "</tbody></table>" if fab else "")
            + ("<p class='note'>DERIVED — computed from verbatim source figures (grounded; "
               "confirm the arithmetic):</p><table><thead><tr><th>Kind</th><th>Draft value"
               "</th><th>Derivation</th></tr></thead><tbody>" + der + "</tbody></table>" if der else "")
            + ("<p class='note'>AMBER — in the source contract, absent from the draft "
               "(8-K disclosure is selective; confirm none is legally material):</p>"
               "<table><thead><tr><th>Kind</th><th>Source value</th><th>Context"
               "</th></tr></thead><tbody>" + om + "</tbody></table>" if om else ""))
    return f"""<!doctype html><html><head><meta charset="utf-8"><style>
      body {{ font-family: 'Times New Roman', Georgia, serif; color: #000; padding: 50px 60px;
              font-size: 13px; line-height: 1.45; }}
      h1 {{ font-size: 18px; }} h2 {{ font-size: 14px; margin-top: 22px; }}
      table {{ border-collapse: collapse; width: 100%; margin-top: 8px; font-size: 12px; }}
      td, th {{ border: 1px solid #999; padding: 5px 8px; text-align: left; vertical-align: top; }}
      th {{ background: #1a2238; color: #fff; }}
      tr.unverified td {{ color: #b4232a; font-weight: 600; }}
      .note {{ color: #555; font-size: 12px; }}
      .verdict {{ font-weight: 700; padding: 8px 10px; border-radius: 4px; margin-top: 8px; }}
      .gr-red {{ background: #fbe6e7; color: #8f1a1f; }}
      .gr-amber {{ background: #fdf1dc; color: #7a5200; }}
      .gr-green {{ background: #e6f4ea; color: #1e6b34; }}
    </style></head><body>
      <h1>8-K Draft — Review Materials (Item {esc(draft.get('item',''))})</h1>
      <p class="note">Companion to the drafted 8-K for source contract
      {esc(draft.get('_source_contract','—'))}. NOT part of the filing — for internal
      review only. Verify every fact below against the source contract.</p>
      {guard_html}
      {"<h2>SEC requirement checks</h2><table><thead><tr><th>Status</th><th>Requirement</th></tr></thead><tbody>" + compliance_rows + "</tbody></table>" if compliance_rows else ""}
      {"<h2>Precedents used (style reference only)</h2><ul>" + precedents + "</ul>" if precedents else ""}
      {"<h2>Fact -&gt; source trace</h2><table><thead><tr><th>Fact</th><th>Source quote</th><th>Verified</th></tr></thead><tbody>" + facts_rows + "</tbody></table>" if facts_rows else ""}
      {"<h2>All terms extracted from the contract</h2><p class='note'>The disclosure states only the material terms, per 8-K convention. This is the full set the review engine extracted &mdash; use it to confirm nothing material was left out.</p><table><thead><tr><th>Term</th><th>Value</th></tr></thead><tbody>" + all_terms_rows + "</tbody></table>" if all_terms_rows else ""}
    </body></html>"""


def _render_pdf(html: str) -> bytes:
    """Render HTML to PDF via a local, offscreen Chromium (Playwright) — no
    network access, same engine used to convert reference filings."""
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.set_content(html, wait_until="load")
        pdf_bytes = page.pdf(format="Letter", margin={"top": "0.6in", "bottom": "0.6in",
                                                        "left": "0.6in", "right": "0.6in"})
        browser.close()
    return pdf_bytes


def draft_to_pdf(draft: dict) -> bytes:
    """The clean filing PDF (no review material)."""
    return _render_pdf(_draft_html(draft))


def review_to_pdf(draft: dict) -> bytes:
    """The review-pack PDF (separate from the filing)."""
    return _render_pdf(_review_html(draft))
